"""
Gerador DAM Word — padrão Cast Group
Produz o Documento de Arquitetura de Melhoria no template oficial
"""
import io
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

AZUL     = RGBColor(0x1F, 0x4E, 0x79)
AZUL_CL  = RGBColor(0x2E, 0x75, 0xB6)
CINZA_HD = RGBColor(0xD9, 0xE1, 0xF2)
BRANCO   = RGBColor(0xFF, 0xFF, 0xFF)

def set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color)
    shd.set(qn('w:val'),  'clear')
    tcPr.append(shd)

def h1(doc, text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = AZUL
    p.runs[0].font.size = Pt(14)
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after  = Pt(8)
    return p

def h2(doc, text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = AZUL_CL
    p.runs[0].font.size = Pt(12)
    p.paragraph_format.space_before = Pt(10)
    return p

def body(doc, text):
    p = doc.add_paragraph(text)
    if p.runs:
        p.runs[0].font.size = Pt(11)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_table_row(table, col1, col2, bg=None, bold_col1=True):
    row = table.add_row()
    c1, c2 = row.cells[0], row.cells[1]
    c1.text = col1
    c2.text = col2
    if bold_col1:
        c1.paragraphs[0].runs[0].bold = True
    if bg:
        set_cell_bg(c1, bg)
        set_cell_bg(c2, bg)
    c1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    c2.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    return row

def generate_dam(sections: dict, payload: dict) -> io.BytesIO:
    doc = Document()

    # Margens
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3)
        section.right_margin  = Cm(2.5)

    # Fonte padrão
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

    data_hoje = datetime.date.today().strftime("%d/%m/%Y")
    titulo_demanda = sections.get("titulo", "Proposta SAP")[:80]

    # ══ CAPA ══
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("CAST GROUP PARTNER")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = AZUL

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(titulo_demanda)
    run.bold = True
    run.font.size = Pt(14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("DAM — Documento de Arquitetura de Melhoria — v1")
    run.font.size = Pt(12)
    run.font.color.rgb = AZUL_CL

    doc.add_paragraph()

    # Tabela capa
    t = doc.add_table(rows=0, cols=2)
    t.style = 'Table Grid'
    t.columns[0].width = Inches(2.0)
    t.columns[1].width = Inches(4.5)

    campos = [
        ("Dados da Solicitação", "", "D9E1F2"),
        ("Nome do Cliente",      "",  None),
        ("Solicitante",          "",  None),
        ("", "", None),
        ("Preenchimento Cast Group", "", "D9E1F2"),
        ("Data criação",         data_hoje, None),
        ("Arquiteto",            "Sil-Proposta (IA)", None),
        ("Código PMS",           "OP —",    None),
    ]
    for c1, c2, bg in campos:
        add_table_row(t, c1, c2, bg=bg)

    doc.add_page_break()

    # ══ SUMÁRIO ══
    h1(doc, "Sumário")
    body(doc, "1. Necessidade")
    body(doc, "2. Solução")
    body(doc, "3. Premissas Gerais")
    body(doc, "4. Equipe do Projeto")
    body(doc, "5. Análise de Impactos")
    body(doc, "6. Cronograma")
    body(doc, "7. Investimento")
    body(doc, "8. Condições de Faturamento")
    doc.add_page_break()

    # ══ 1. NECESSIDADE ══
    h1(doc, "1. Necessidade")

    ufs = ", ".join(sections.get("ufs", []))
    ver = sections.get("versao_sap", "")
    tipo = sections.get("tipo_projeto", "")

    body(doc, f"O cliente opera SAP {ver} e necessita de desenvolvimento/configuração "
              f"para adequação fiscal e operacional nas UFs: {ufs}. "
              f"Tipo de projeto: {tipo}.")

    h2(doc, "1.1 Benefício esperado pelo cliente")
    t = doc.add_table(rows=0, cols=2)
    t.style = 'Table Grid'
    add_table_row(t, "Benefício esperado pelo cliente",
        "Garantir conformidade fiscal, eliminar riscos de autuação e automatizar processos manuais.",
        bg="D9E1F2")

    h2(doc, "1.2 Quadro resumo do processo atual")
    t = doc.add_table(rows=0, cols=2)
    t.style = 'Table Grid'
    rows_atual = [
        ("Processo atual", sections.get("necessidade", "Processo atual descrito na RFP.")),
        ("Transações do processo atual", "VA01, VF01, J1BNFE, FI (baixa de títulos)"),
        ("Sistemas externos e interfaces", "Gateway de pagamento / TEF / maquininha"),
        ("Volume de dados", "A confirmar com o cliente"),
        ("Áreas e processos impactados", "Vendas, Faturamento, Financeiro (FI), TI"),
        ("Demanda relacionada a compliance?", "Sim — legislação fiscal vigente"),
    ]
    for c1, c2 in rows_atual:
        add_table_row(t, c1, c2)

    h2(doc, "1.3 Quadro resumo do processo futuro")
    t = doc.add_table(rows=0, cols=2)
    t.style = 'Table Grid'
    rows_fut = [
        ("Processo futuro", "Integração automatizada entre os sistemas de pagamento e o SAP, com transmissão fiscal automatizada."),
        ("Entradas do processo", "Dados de pagamento vindos do gateway/TEF/maquininha"),
        ("Saídas do processo", "NF-e com Grupo YA preenchido (Cenário 1) e Evento ECONF (Cenário 2)"),
        ("Clientes", "N/A"),
    ]
    for c1, c2 in rows_fut:
        add_table_row(t, c1, c2)

    doc.add_page_break()

    # ══ 2. SOLUÇÃO ══
    h1(doc, "2. Solução")
    h2(doc, "2.1 Resumo da Solução")

    # Fluxo
    drc_res   = sections.get("plano",{})
    needs_cpi = drc_res.get("needs_cpi", False)
    canal = "SAP CPI (evento sem suporte DRC nativo)" if needs_cpi else "SAP DRC"
    body(doc, f"A solução utiliza {canal} para transmissão fiscal. "
              f"Fluxo: FI Baixa → ABAP Z monta payload → {canal} → SEFAZ → protocolo retorna ao ECC.")

    reforma = sections.get("reforma", {})
    if reforma.get("decisao") == "fazer_agora":
        body(doc, "⚠ Reforma Tributária (LC 214): impacto identificado — entregáveis incluídos no escopo.")
    elif reforma.get("decisao") == "planejar":
        body(doc, "ℹ Reforma Tributária (LC 214): recomenda-se planejamento. Roadmap incluído como seção informativa.")

    h2(doc, "2.2 Escopo da melhoria")
    entregaveis = sections.get("entregaveis", [])
    if not entregaveis:
        entregaveis = [
            "Especificação funcional — análise da legislação e dos cenários",
            "Configuração/especificação FI — conciliação bancária e trigger do evento",
            "Configuração/especificação SD — BAPI, BAdI, RFC, monitor de eventos",
            "Desenvolvimento ABAP — BAPI Z, BAdI, RFC Z, evento, monitor, CPI",
            "Testes de validação — unitários e integrados",
            "Auxílio testes integrados (1 dia útil por frente)",
            "Acompanhamento Go Live (1 dia útil por frente)",
        ]

    t = doc.add_table(rows=0, cols=4)
    t.style = 'Table Grid'
    hdr = t.add_row()
    for i, h in enumerate(["Mód.", "Produto / Entregável", "Escopo", "Premissas"]):
        hdr.cells[i].text = h
        hdr.cells[i].paragraphs[0].runs[0].bold = True
        set_cell_bg(hdr.cells[i], "D9E1F2")

    modulos_rows = [
        ("SD",   "Entendimento do cenário",              "Sim", "Análise da legislação e geração do DAM"),
        ("FI",   "Entendimento do cenário",              "Sim", "Análise do fluxo financeiro"),
        ("FI",   "Configuração / especificação",         "Sim", "Conciliação bancária + trigger do evento"),
        ("SD",   "Configuração / especificação",         "Sim", "BAPI Z + BAdI + RFC + monitor"),
        ("ABAP", "Ajuste de programas",                  "Sim", "BAPI, BAdI, RFC, evento, monitor, CPI"),
        ("SD",   "Testes de validação",                  "Sim", "Testes unitários — interface da maquininha"),
        ("FI",   "Testes de validação",                  "Sim", "Testes unitários — dados para XML e evento"),
        ("ABAP", "Testes de validação",                  "Sim", "Suporte aos testes"),
        ("SD",   "Auxílio testes integrados (1 dia útil)","Sim",""),
        ("FI",   "Auxílio testes integrados (1 dia útil)","Sim",""),
        ("ABAP", "Auxílio testes integrados (1 dia útil)","Sim",""),
        ("SD",   "Acompanhamento Go Live (1 dia útil)",  "Sim", ""),
        ("FI",   "Acompanhamento Go Live (1 dia útil)",  "Sim", ""),
        ("ABAP", "Acompanhamento Go Live (1 dia útil)",  "Sim", ""),
    ]
    for row_data in modulos_rows:
        row = t.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = val

    doc.add_page_break()

    # ══ 3. PREMISSAS ══
    h1(doc, "3. Premissas Gerais")
    premissas = sections.get("premissas", [])
    for i, p_text in enumerate(premissas, 1):
        p = doc.add_paragraph(style='List Number')
        p.add_run(p_text)

    # Definição de alteração de escopo
    h2(doc, "Definição de Alteração de Escopo")
    for item in [
        "Qualquer inclusão de novos itens ou alteração de já existentes não previstos nesta proposta.",
        "Alteração de regras de negócio, lógicas, layout ou apresentações previamente acordadas.",
        "Alteração de cronogramas sem aviso mínimo de 1 semana.",
        "Liberação de consultores sem aviso prévio de 7 dias.",
        "Quaisquer atividades não previstas neste documento.",
    ]:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(item)

    doc.add_page_break()

    # ══ 4. EQUIPE ══
    h1(doc, "4. Equipe do Projeto")
    equipe = sections.get("equipe", [])
    for recurso in equipe:
        frente = recurso.get("frente","")
        nivel  = recurso.get("nivel","Senior")
        body(doc, f"• Consultor {frente} — {nivel}")

    doc.add_page_break()

    # ══ 5. ANÁLISE DE IMPACTOS ══
    h1(doc, "5. Análise de Impactos")
    t = doc.add_table(rows=0, cols=5)
    t.style = 'Table Grid'
    hdr = t.add_row()
    for i, h in enumerate(["Id","Impacto","Solução de Contorno","Probabilidade","Classificação"]):
        hdr.cells[i].text = h
        hdr.cells[i].paragraphs[0].runs[0].bold = True
        set_cell_bg(hdr.cells[i], "D9E1F2")
    row = t.add_row()
    vals = ["01","Erros durante o Go Live",
            "Extenso ciclo de testes em QAS + homologação SEFAZ",
            "Média","Moderado"]
    for i, v in enumerate(vals):
        row.cells[i].text = v

    # ══ 6. CRONOGRAMA ══
    doc.add_page_break()
    h1(doc, "6. Cronograma")
    equipe_r = sections.get("equipe", [])
    semanas = max((r.get("dias",10)//5) for r in equipe_r) if equipe_r else 4
    body(doc, f"Macro cronograma: aproximadamente {semanas} semanas para execução das atividades e suporte.")
    body(doc, "Este cronograma será detalhado após a aprovação desta proposta e poderá sofrer alterações.")
    body(doc, "O início do projeto será planejado a partir da aprovação desta proposta.")

    # ══ 7. INVESTIMENTO ══
    doc.add_page_break()
    h1(doc, "7. Investimento")
    total_h = sections.get("total_horas", 528)
    comercial = sections.get("comercial", {})
    valor = comercial.get("valor_referencia", 0) or (total_h * 230)

    t = doc.add_table(rows=0, cols=2)
    t.style = 'Table Grid'
    add_table_row(t, "Título da Demanda", "Total", bg="D9E1F2")
    add_table_row(t, titulo_demanda[:60],
        f"R$ {valor:,.2f}".replace(",","X").replace(".",",").replace("X","."))

    body(doc, "")
    body(doc, "É importante destacar que, conforme o modelo de Sustentação da Cast Group, "
              "todas as melhorias são consideradas como 'valor fechado' e possuem garantia de 30 dias.")

    # ══ 8. CONDIÇÕES DE FATURAMENTO ══
    h1(doc, "8. Condições de Faturamento")
    for item in [
        "Esta proposta tem validade de 30 (trinta) dias.",
        "Os valores incluem ISS, PIS e COFINS atualmente em vigor.",
        "Faturamento: 50% na aprovação desta proposta + 50% no Go-Live.",
        "Em caso de paralisação: Cast reserva-se o direito de faturar o % de avanço.",
        "Garantia pós go-live: 30 dias.",
    ]:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(item)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
